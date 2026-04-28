from docx import Document
from docx.shared import Inches, Pt
import os

def create_docx(input_md, output_docx):
    doc = Document()
    
    # Title
    doc.add_heading('Master Engineering Specification: QR-ZITIS', 0)
    
    with open(input_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif '|' in line and '--' not in line:
            # Simple table handling
            p = doc.add_paragraph(style='No Spacing')
            run = p.add_run(line.replace('|', ' | '))
            run.font.size = Pt(10)
            run.italic = True
        elif line.startswith('```'):
            continue # Skip code block markers
        else:
            doc.add_paragraph(line)

    doc.save(output_docx)
    print(f"Word document successfully generated at: {output_docx}")

if __name__ == "__main__":
    input_file = "SENIOR_ENGINEER_PRD.md"
    output_file = "QR_ZITIS_Master_Engineering_PRD.docx"
    create_docx(input_file, output_file)
